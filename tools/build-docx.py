"""Build a DOCX document from a documentation directory.

Usage:
    python tools/build-docx.py docs/system-design

Reads sidebar.md to determine file order, renders diagrams to images,
and produces a single DOCX file in output/.
"""

import argparse
import logging
import re
import shutil
import subprocess
import sys
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt


def parse_sidebar(sidebar_path):
    """Parse sidebar.md and return ordered list of relative markdown paths."""
    files = []
    with open(sidebar_path, "r", encoding="utf-8") as handle:
        for line in handle:
            match = re.search(r"\[.*?\]\((.*?)\)", line)
            if match:
                link = match.group(1)
                if not link or link.startswith("http"):
                    continue
                if not link.endswith(".md"):
                    link += ".md"
                files.append(link)
    return files


def render_mermaid(code, output_path):
    """Render Mermaid code to a PNG image using mermaidx."""
    try:
        import mermaidx
        diagram = mermaidx.render(code)
        diagram.save(output_path)
        return True
    except Exception as error:
        logging.warning("WARNING: Mermaid rendering failed: %s", error)
        return False


def render_graphviz(code, output_path):
    """Render Graphviz code to a PNG image using dot (optional system dependency)."""
    if not shutil.which("dot"):
        logging.warning("WARNING: dot not found, skipping Graphviz diagram")
        return False
    try:
        subprocess.run(
            ["dot", "-Tpng", "-o", output_path],
            input=code.encode(),
            check=True,
            capture_output=True,
        )
        return True
    except subprocess.CalledProcessError as error:
        logging.warning("WARNING: Graphviz rendering failed: %s", error.stderr.decode())
        return False


def parse_markdown_blocks(content):
    """Parse markdown content into blocks (headings, paragraphs, code, images, diagrams)."""
    blocks = []
    lines = content.split("\n")
    index = 0

    while index < len(lines):
        line = lines[index]

        # Setext heading level 1 (next line is ===)
        if index + 1 < len(lines) and re.match(r"^=+\s*$", lines[index + 1]):
            blocks.append({"type": "heading", "level": 1, "text": line.strip()})
            index += 2
            continue

        # Setext heading level 2 (next line is ---)
        if index + 1 < len(lines) and re.match(r"^-{3,}\s*$", lines[index + 1]):
            blocks.append({"type": "heading", "level": 2, "text": line.strip()})
            index += 2
            continue

        # ATX headings
        heading_match = re.match(r"^(#{1,6})\s+(.+)", line)
        if heading_match:
            level = len(heading_match.group(1))
            blocks.append({"type": "heading", "level": level, "text": heading_match.group(2).strip()})
            index += 1
            continue

        # Fenced code block
        fence_match = re.match(r"^```(\w*)", line)
        if fence_match:
            language = fence_match.group(1)
            code_lines = []
            index += 1
            while index < len(lines) and not re.match(r"^```\s*$", lines[index]):
                code_lines.append(lines[index])
                index += 1
            index += 1  # skip closing fence
            code_content = "\n".join(code_lines)

            if language == "mermaid":
                blocks.append({"type": "mermaid", "code": code_content})
            elif language in ("graphviz", "dot"):
                blocks.append({"type": "graphviz", "code": code_content})
            else:
                blocks.append({"type": "code", "language": language, "code": code_content})
            continue

        # Image
        image_match = re.match(r"!\[([^\]]*)\]\(([^)]+)\)", line)
        if image_match:
            blocks.append({"type": "image", "alt": image_match.group(1), "path": image_match.group(2)})
            index += 1
            continue

        # Empty line
        if not line.strip():
            index += 1
            continue

        # Regular paragraph (collect consecutive non-empty lines)
        paragraph_lines = []
        while index < len(lines) and lines[index].strip() and not re.match(r"^(#{1,6}\s|```|!\[)", lines[index]):
            # Check if next line is a setext underline
            if index + 1 < len(lines) and re.match(r"^[=-]{3,}\s*$", lines[index + 1]):
                break
            paragraph_lines.append(lines[index])
            index += 1
        if paragraph_lines:
            blocks.append({"type": "paragraph", "text": " ".join(paragraph_lines)})
        continue

    return blocks


def add_heading(document, text, level):
    """Add a heading to the document."""
    document.add_heading(text, level=level)


def add_paragraph(document, text):
    """Add a paragraph with basic inline formatting."""
    paragraph = document.add_paragraph()
    parts = re.split(r"(\*\*.*?\*\*|__.*?__|_.*?_|\*.*?\*|`.*?`)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("__") and part.endswith("__"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("_") and part.endswith("_"):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        elif part.startswith("*") and part.endswith("*"):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Courier New"
            run.font.size = Pt(9)
        else:
            paragraph.add_run(part)


def add_code_block(document, code):
    """Add a code block as a formatted paragraph."""
    paragraph = document.add_paragraph()
    paragraph.style = document.styles["No Spacing"]
    run = paragraph.add_run(code)
    run.font.name = "Courier New"
    run.font.size = Pt(9)


def add_image(document, image_path, docs_dir):
    """Add an image to the document."""
    full_path = Path(docs_dir) / image_path
    if full_path.exists():
        document.add_picture(str(full_path), width=Inches(5.5))
    else:
        logging.warning("WARNING: Image not found: %s", image_path)


def find_sidebar(docs_path):
    """Find sidebar.md in the given directory or its parent."""
    sidebar_path = docs_path / "sidebar.md"
    if sidebar_path.exists():
        return sidebar_path
    parent_sidebar = docs_path.parent / "sidebar.md"
    if parent_sidebar.exists():
        return parent_sidebar
    return None


def build_docx(docs_dir, output_dir):
    """Build a DOCX document from the documentation directory."""
    docs_path = Path(docs_dir)
    sidebar_path = find_sidebar(docs_path)

    if not sidebar_path:
        logging.error("ERROR: sidebar.md not found in %s or its parent", docs_path)
        sys.exit(1)

    # Parse sidebar for file order, resolve relative to sidebar location
    raw_files = parse_sidebar(sidebar_path)
    sidebar_dir = sidebar_path.parent
    file_list = []
    for file_path in raw_files:
        full_path = sidebar_dir / file_path
        if full_path.resolve().is_relative_to(docs_path.resolve()):
            file_list.append(str(full_path.relative_to(docs_path)))

    logging.info("Found %d files in sidebar.md for %s", len(file_list), docs_path.name)

    # Create output directory
    output_path = Path(output_dir)
    output_path.mkdir(exist_ok=True)

    # Create temporary directory for rendered diagrams
    diagrams_dir = output_path / "diagrams"
    diagrams_dir.mkdir(exist_ok=True)

    # Create document
    template_path = Path("tools/reference.docx")
    if template_path.exists():
        document = Document(str(template_path))
    else:
        document = Document()

    diagram_counter = 0

    for file_path in file_list:
        full_path = docs_path / file_path
        if not full_path.exists():
            logging.warning("WARNING: Skipping %s (not found)", file_path)
            continue

        logging.info("  Processing %s", file_path)
        content = full_path.read_text(encoding="utf-8")
        if not content.strip():
            continue

        blocks = parse_markdown_blocks(content)

        for block in blocks:
            if block["type"] == "heading":
                add_heading(document, block["text"], block["level"])
            elif block["type"] == "paragraph":
                add_paragraph(document, block["text"])
            elif block["type"] == "code":
                add_code_block(document, block["code"])
            elif block["type"] == "mermaid":
                diagram_counter += 1
                image_path = diagrams_dir / f"mermaid_{diagram_counter}.png"
                if render_mermaid(block["code"], str(image_path)):
                    document.add_picture(str(image_path), width=Inches(5.5))
                else:
                    add_code_block(document, block["code"])
            elif block["type"] == "graphviz":
                diagram_counter += 1
                image_path = diagrams_dir / f"graphviz_{diagram_counter}.png"
                if render_graphviz(block["code"], str(image_path)):
                    document.add_picture(str(image_path), width=Inches(5.5))
                else:
                    add_code_block(document, block["code"])
            elif block["type"] == "image":
                add_image(document, block["path"], docs_path)

    # Determine output filename from directory name
    output_name = docs_path.name
    output_file = output_path / f"{output_name}.docx"
    document.save(str(output_file))
    logging.info("Saved: %s", output_file)


def parse_args():
    """Parse command line arguments."""
    parser = argparse.ArgumentParser(
        description="Build a DOCX document from a documentation directory.")
    parser.add_argument(
        "docs_dir",
        help="Path to the documentation directory (e.g., docs/system-design)")
    parser.add_argument(
        "-o", "--output",
        default="output",
        help="Output directory (default: output)")
    parser.add_argument(
        "-v", "--verbose",
        action="store_true",
        help="Enable verbose output")
    return parser.parse_args()


def main():
    """Entry point."""
    args = parse_args()
    logging.basicConfig(
        level=logging.DEBUG if args.verbose else logging.INFO,
        format="%(message)s")
    build_docx(args.docs_dir, args.output)


if __name__ == "__main__":
    main()
